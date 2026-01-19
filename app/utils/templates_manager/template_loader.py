# Lógica para cargar y gestionar plantillas de Excel y Word

from openpyxl import load_workbook
from docx import Document
import os

class TemplateLoader:
    def __init__(self, template_dir=None):
        if template_dir is None:
            template_dir = os.path.join(os.path.dirname(__file__), '..', '..', '..', 'template')
        self.template_dir = template_dir

    def load_template(self, template_path):
        full_path = os.path.join(self.template_dir, template_path)
        if template_path.endswith('.xlsx'):
            return load_workbook(full_path)
        elif template_path.endswith('.xls'):
            return load_workbook(full_path, engine='xlrd')
        elif template_path.endswith('.docx'):
            return Document(full_path)
        else:
            raise ValueError("Unsupported template format")

    def fill_template(self, template, data):
        # Mapeo de placeholders a claves de data
        mapping = {
            'cuidad': 'cuidad_colaborador',
            'direccion': 'direccion',
            'fecha': 'fecha_entrega_documentos',  # Campo oculto con fecha formateada
            'colaborador_apellidos': 'colaborador_apellidos',
            'colaborador_nombres': 'colaborador_nombres',
            'cargo': 'cargo',
            'colaborador': 'colaborador',
            'responsable_DTH': 'responsable_DTH',
            'departamento': 'area_departamento',
            'reemplazo': 'reemplazo',
            'cedula': 'cedula',
            'lugar_trabajo': 'lugar_trabajo',
            'recibido': 'recibido_por',
            'analisis_talento': 'analista_th',
            'trabajadora_social': 'trabajadora_social',
            'asistente_remuneracion': 'remuneraciones',
            'analista_remuneracion': 'remuneraciones',
            'jefes_seguros': 'seguros',
            'supervisor': 'seguridad',
            'riesgo': 'riesgo',
            'coordinador': 'coordinador_calidad',
            'ingeniero_calidad': 'ingeniero_calidad',
            'medico_ocupacional': 'departamento_medico',
            'enfermera': 'enfermera',  # Si no existe, será vacío
            'ingeniero_electrico': 'ingeniero_electrico',
            'superintendente': 'superintendente',
            'dia': 'dia',
            'mes': 'mes',
            'anio': 'anio'
        }
        
        if isinstance(template, load_workbook):
            # Llenar Excel - todas las hojas
            for sheet in template.worksheets:
                for placeholder, key in mapping.items():
                    value = data.get(key, '')
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                cell.value = cell.value.replace('{{' + placeholder + '}}', str(value))
            return template
        elif isinstance(template, Document):
            # Llenar Word - párrafos principales
            for paragraph in template.paragraphs:
                for placeholder, key in mapping.items():
                    value = data.get(key, '')
                    if '{{' + placeholder + '}}' in paragraph.text:
                        paragraph.text = paragraph.text.replace('{{' + placeholder + '}}', str(value))
            # Llenar tablas
            for table in template.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, key in mapping.items():
                                value = data.get(key, '')
                                if '{{' + placeholder + '}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{' + placeholder + '}}', str(value))
            return template
        else:
            raise ValueError("Unsupported template type")

    def save_filled_template(self, filled_template, output_path):
        if isinstance(filled_template, load_workbook):
            filled_template.save(output_path)
        elif isinstance(filled_template, Document):
            filled_template.save(output_path)
